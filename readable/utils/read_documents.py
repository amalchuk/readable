from logging import Logger
from logging import getLogger as get_logger
from pathlib import Path as _P
from typing import Callable, Final, Iterator, Optional

from docx import Document as DOCXDocument
from fitz import Document as PDFDocument
from fitz import TEXT_INHIBIT_SPACES

__all__: Final[list[str]] = ["microsoft_word_document", "pdf_document", "read_document", "text_document"]

logger: Logger = get_logger(__name__)


def microsoft_word_document(filename: _P, /) -> str:
    """
    Extract the text from an existing ``docx`` file.
    """
    def wrapper() -> Iterator[str]:
        document = DOCXDocument(filename)
        for paragraph in document.paragraphs:
            yield paragraph.text.strip()

    return "\n".join(wrapper())


def pdf_document(filename: _P, /) -> str:
    """
    Extract the text from an existing ``pdf`` file.
    """
    def wrapper() -> Iterator[str]:
        document = PDFDocument(filename)
        for page in document:
            yield page.getText("text", flags=TEXT_INHIBIT_SPACES).strip()

    return "".join(wrapper())


def text_document(filename: _P, /) -> str:
    """
    Extract the text from an existing ``plain`` file.
    """
    def wrapper() -> Iterator[str]:
        with filename.open(encoding="utf-8") as istream:
            while text := istream.read(4096):
                yield text.strip()

    return "".join(wrapper())


def read_document(filename: _P, /) -> Optional[str]:
    """
    Extract the text from an existing ``docx``, ``pdf`` or ``plain`` file.
    """
    allowed_functions: dict[str, Callable[[_P], str]] = {
        ".docx": microsoft_word_document,
        ".pdf": pdf_document,
        ".txt": text_document
    }
    try:
        extension: str = filename.suffix.lower()
        callback: Callable[[_P], str] = allowed_functions[extension]
        return callback(filename)

    except Exception as exception:
        logger.exception(exception)
        return None
