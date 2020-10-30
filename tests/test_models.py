from io import BytesIO

from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.db.models.signals import post_save
from django.test.testcases import TestCase
from docx import Document as DOCXDocument

from readable.models import Documents
from readable.models import Staff
from readable.utils.signals import documents_uploaded


class TestDocuments(TestCase):
    def setUp(self) -> None:
        document = DOCXDocument()
        document.add_heading("Document Title", 0)
        document.add_paragraph("Lorem ipsum dolor sit amet, consectetur adipiscing elit.")

        docx_document = BytesIO()
        document.save(docx_document)
        self.docx_document = ContentFile(docx_document.getvalue(), "document.docx")

        user = User.objects.create_user(username="staff", password="staff")
        self.staff = Staff.objects.create(user=user)

    def test_upload_directory(self) -> None:
        # Temporarily disable the "documents_uploaded" signal:
        post_save.disconnect(documents_uploaded, Documents)

        document = Documents.objects.create(filename=self.docx_document, uploaded_by=self.staff)
        self.assertEqual(document.realname, self.docx_document.name)
        self.assertEqual(document.filename, f"{document.id!s}{document.path.suffix}")

        self.assertTrue(document.unavailable)
        document.status = Documents.Status.FINISHED
        document.save(update_fields=["status"])
        self.assertFalse(document.unavailable)
