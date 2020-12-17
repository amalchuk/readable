import logging
from os import urandom
from pathlib import Path as _P
from typing import List

from django.test.testcases import SimpleTestCase as TestCase
from docx import Document as DOCXDocument
from reportlab.pdfgen.canvas import Canvas as PDFDocument

from readable.utils.read_documents import microsoft_word_document
from readable.utils.read_documents import pdf_document
from readable.utils.read_documents import read_document
from readable.utils.read_documents import text_document
from readable.utils.temporary import temporary_directory


class TestReadDocuments(TestCase):
    def setUp(self) -> None:
        logging.disable(logging.ERROR)
        self.lorem: List[str] = [
            "Lorem ipsum dolor sit amet, consectetur adipiscing elit.",
            "Aliquam sollicitudin lorem non hendrerit porttitor.",
            "Donec tincidunt quam ut nibh varius pharetra.",
            "Praesent venenatis metus a elit ullamcorper cursus.",
            "Etiam ullamcorper quam quis turpis accumsan, consequat maximus risus tristique."
        ]
        self.temporary_directory = temporary_directory()

    def test_microsoft_word_document(self) -> None:
        tempfile = str(self.temporary_directory / "document01.docx")

        document = DOCXDocument()
        for paragraph in self.lorem:
            document.add_paragraph(paragraph)

        document.save(tempfile)
        self.assertEqual(microsoft_word_document(_P(tempfile)), "\n".join(self.lorem))

    def test_pdf_document(self) -> None:
        tempfile = str(self.temporary_directory / "document01.pdf")

        document = PDFDocument(tempfile)
        text_object = document.beginText(100, 100)

        for paragraph in self.lorem:
            text_object.textLine(paragraph)

        document.drawText(text_object)
        document.showPage()
        document.save()
        self.assertEqual(pdf_document(_P(tempfile)), "\n".join(self.lorem))

    def test_text_document(self) -> None:
        tempfile = str(self.temporary_directory / "document01.txt")
        document = _P(tempfile)
        document.write_text("\n".join(self.lorem))
        self.assertEqual(text_document(_P(tempfile)), "\n".join(self.lorem))

    def test_read_document(self) -> None:
        tempfile = str(self.temporary_directory / "document02.txt")
        document = _P(tempfile)
        document.write_text("\n".join(self.lorem))

        text = read_document(_P(tempfile))
        self.assertIsNotNone(text)
        self.assertEqual(text, "\n".join(self.lorem))

        tempfile = str(self.temporary_directory / "document.bin")
        document = _P(tempfile)
        document.write_bytes(urandom(128))
        self.assertIsNone(read_document(document))

    def tearDown(self) -> None:
        logging.disable(logging.NOTSET)
